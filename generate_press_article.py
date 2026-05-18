"""
Generate a press article docx for the EDGE Flowering Plants × Flora of the World project.
"""

import csv
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── helpers ────────────────────────────────────────────────────────────────────

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # underline + blue colour
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E79")
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_paragraph(doc, text, style="Normal", bold=False, italic=False,
                  size=None, color=None, space_before=None, space_after=None,
                  align=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def italic_species(paragraph, text):
    """Add a run with italic species name embedded in plain text.
    text can contain *...*  markers for italic spans."""
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            if part:
                paragraph.add_run(part)


def threat_label(code):
    mapping = {
        "CR": "Critically Endangered",
        "EN": "Endangered",
        "VU": "Vulnerable",
        "NT": "Near Threatened",
        "LC": "Least Concern",
        "EW": "Extinct in the Wild",
        "EX": "Extinct",
        "thr": "Threatened (predicted)",
        "not": "Not threatened (predicted)",
    }
    return mapping.get(code, code)


def threat_color(code):
    """Return a hex fill colour for threat status."""
    return {
        "CR": "C00000",
        "EN": "E26B0A",
        "VU": "F4B942",
        "NT": "92D050",
        "LC": "00B050",
        "EW": "7030A0",
        "thr": "FF9999",
        "not": "D9D9D9",
    }.get(code, "FFFFFF")


# ── load data ──────────────────────────────────────────────────────────────────

MATCHED_CSV = "/Users/sven/Documents/Current_projects/Endowed_Chair/EDGE_flowering_plants/EDGE_FotW_matched_species.csv"

matched = []
with open(MATCHED_CSV) as f:
    for row in csv.DictReader(f):
        matched.append(row)

edge_list = [r for r in matched if r.get("EDGE_List", "").strip().lower() in ("y", "yes")]
edge_list_sorted = sorted(edge_list, key=lambda r: int(r["EDGE_rank"]))

ew_edge = [r for r in edge_list if r.get("threat", "").strip() == "EW"]
ew_edge_sorted = sorted(ew_edge, key=lambda r: int(r["EDGE_rank"]))

top25 = edge_list_sorted[:25]

total_ed_fotw = sum(float(r["ed_med"]) for r in matched)
total_ed_all = 1_026_781.1   # from report
ed_pct = total_ed_fotw / total_ed_all * 100

n_edge_list = 9945
n_matched = len(edge_list)

# ── build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── TITLE ──────────────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("Boise State Scientist Co-Authors Landmark Study\non the World's Most Irreplaceable Plants")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
p_title.paragraph_format.space_after = Pt(4)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run(
    "Prof. Sven Buerki, Davidson Endowed Chair of Botany at Boise State University, contributes to a "
    "Science paper scoring the evolutionary uniqueness of all flowering plants — and leads the effort "
    "to document 229 priority species in Flora of the World."
)
run2.italic = True
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
p_sub.paragraph_format.space_after = Pt(14)

# ── KEY STATISTICS BOX ────────────────────────────────────────────────────────
doc.add_paragraph()
p_box_title = doc.add_paragraph()
run = p_box_title.add_run("  Key Statistics at a Glance")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
p_box_title.paragraph_format.space_before = Pt(0)
p_box_title.paragraph_format.space_after  = Pt(0)
# Blue background via table trick
stats_table = doc.add_table(rows=1, cols=1)
stats_table.style = "Table Grid"
cell = stats_table.cell(0, 0)
set_cell_bg(cell, "1F4E79")
cell.width = Inches(6.5)
cp = cell.paragraphs[0]
r = cp.add_run("Key Statistics at a Glance")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
cp.paragraph_format.space_before = Pt(4)
cp.paragraph_format.space_after  = Pt(4)

stats_data = [
    ("Flowering plant species scored for EDGE",        "335,497"),
    ("EDGE.List priority species (globally threatened + evolutionarily distinctive)",
                                                       "9,945"),
    ("EDGE.List species documented in Flora of the World",
                                                       f"{n_matched} (2.3% of EDGE.List)"),
    ("Highest EDGE rank documented in FotW",           "#2 — Amorphophallus lewallei (CR)"),
    ("Evolutionary history (ED) captured in FotW",    f"{total_ed_fotw:,.1f} Myr ({ed_pct:.1f}% of total angiosperm ED)"),
    ("Most evolutionarily distinct species in FotW",   "Amborella trichopoda — 139 million years"),
    ("Extinct-in-the-Wild EDGE species in FotW",       f"{len(ew_edge)} species"),
]
stats_tbl = doc.add_table(rows=len(stats_data), cols=2)
stats_tbl.style = "Table Grid"
for i, (label, value) in enumerate(stats_data):
    row = stats_tbl.rows[i]
    lc = row.cells[0]
    vc = row.cells[1]
    bg = "D9E2F3" if i % 2 == 0 else "EBF0FB"
    set_cell_bg(lc, bg)
    set_cell_bg(vc, bg)
    lp = lc.paragraphs[0]
    lr = lp.add_run(label)
    lr.font.size = Pt(10)
    lr.bold = True
    vp = vc.paragraphs[0]
    vr = vp.add_run(value)
    vr.font.size = Pt(10)
    vr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    lp.paragraph_format.space_before = Pt(2)
    lp.paragraph_format.space_after  = Pt(2)
    vp.paragraph_format.space_before = Pt(2)
    vp.paragraph_format.space_after  = Pt(2)

doc.add_paragraph()

# ── BODY TEXT ────────────────────────────────────────────────────────────────

def section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body_para(doc, text, space_after=6):
    """Add a body paragraph with optional *italic* species markers."""
    p = doc.add_paragraph()
    italic_species(p, text)
    for run in p.runs:
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def pull_quote(doc, quote_text, attribution):
    """Add a styled pull-quote with left border and attribution."""
    # Quote text
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    # Left border via paragraph XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), "1F4E79")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(f"\u201c{quote_text}\u201d")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    # Attribution
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent  = Cm(1.0)
    pa.paragraph_format.space_before = Pt(0)
    pa.paragraph_format.space_after  = Pt(10)
    ra = pa.add_run(f"— {attribution}")
    ra.bold = True
    ra.font.size = Pt(10)
    ra.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    return p


section_heading(doc, "A Landmark Publication in Science")

body_para(doc,
    "A new study published in Science by Forest et al. (2026) provides the first comprehensive assessment "
    "of Evolutionarily Distinct and Globally Endangered (EDGE) scores for all 335,497 known flowering plant "
    "species. The work, a collaboration between the Royal Botanic Gardens, Kew, ZSL (Zoological Society of "
    "London), and partners worldwide, reveals that more than a fifth (21%) of flowering plant evolutionary "
    "history is currently at risk of extinction — nearly double that of jawed vertebrates."
)

body_para(doc,
    "Among the contributors to this international effort is Prof. Sven Buerki, Davidson Endowed Chair of Botany "
    "at Boise State University (BSU). Buerki was a member of Dr. Félix Forest's research group at the Royal "
    "Botanic Gardens, Kew, where he helped lay the foundations for this work. Now leading BSU's Department of "
    "Biological Sciences plant research programme, he brings this global collaboration directly to Idaho — and "
    "to the Flora of the World platform, where the study's findings are being translated into accessible, "
    "open conservation data."
)

body_para(doc,
    "From the full dataset, 9,945 species form the EDGE.List: plants that combine high evolutionary "
    "distinctiveness with serious extinction risk — the world's most irreplaceable species. BSU's role in "
    "connecting this priority list to Flora of the World (FotW) is a direct expression of the conservation "
    "mission funded by the Davidson Endowment."
)

section_heading(doc, "From Gymnosperms to Flowering Plants: A Decade of Collaboration")

body_para(doc,
    "Buerki's involvement in the EDGE programme predates this Science paper. In 2018, he was last author on "
    "a foundational study published in Scientific Reports that produced the first EDGE scores for gymnosperms "
    "— the conifers, cycads, and their relatives — in collaboration with Forest and colleagues at Kew and ZSL. "
    "That paper identified the most evolutionarily distinct and threatened seed plants outside the flowering "
    "plants, and established the analytical framework that has since been applied, at far greater scale, to "
    "all 335,497 angiosperms in the current work."
)

body_para(doc,
    "The progression from gymnosperms to flowering plants reflects a decade-long international partnership "
    "between Boise State University and the Royal Botanic Gardens, Kew — one of the world's leading plant "
    "science institutions. That partnership is now producing tangible conservation outcomes through the "
    "Flora of the World platform."
)

section_heading(doc, "What Are EDGE Scores and Evolutionary Distinctiveness?")

body_para(doc,
    "Evolutionary Distinctiveness (ED) measures how isolated a species is on the tree of life — how much unique "
    "evolutionary history it represents, expressed in millions of years (Myr). A species with high ED has few "
    "close relatives and has been evolving independently for an exceptionally long time. Losing it means losing "
    "an irreplaceable branch of the tree of life."
)

body_para(doc,
    "The EDGE score combines ED with extinction risk (from the IUCN Red List), producing a single metric that "
    "identifies the species most urgently in need of conservation attention. A high EDGE score signals both "
    "extraordinary uniqueness and imminent danger."
)

body_para(doc,
    "A striking example is *Amborella trichopoda* — a small shrub from New Caledonia whose lineage diverged from "
    "all other flowering plants approximately 130 million years ago. With an ED score of 139 Myr, it is the most "
    "evolutionarily distinct angiosperm in FotW, and its documentation on the platform ensures that scientists "
    "and conservationists can access its occurrence data alongside its evolutionary significance."
)

pull_quote(doc,
    "More than two in five plant species are estimated to be threatened with extinction, yet fewer than 20% have "
    "formal threat assessments — compared to over 80% of vertebrate groups. This makes it extremely difficult to "
    "prioritise plants for conservation. Our study provides a powerful framework to rank species for conservation "
    "action by combining their unique evolutionary history with extinction risk. It is the first time this approach "
    "has been applied to a group of this scale, and it will pave the way for other megadiverse groups such as insects.",
    "Prof. Sven Buerki, Davidson Endowed Chair of Botany, Boise State University"
)

section_heading(doc, "Flora of the World Documents 229 Priority EDGE Species")

body_para(doc,
    "A key contribution of Buerki's work at BSU has been to cross-reference the EDGE dataset against Flora of "
    f"the World's global occurrence database. The analysis reveals that {n_matched} of the 9,945 EDGE.List "
    "species (2.3%) are already documented in FotW — with specimen and observation records pinpointing exactly "
    "where these plants have been collected in the wild. These 229 species span remarkable diversity, from the "
    "Critically Endangered *Amorphophallus lewallei* of the Congo Basin (ranked #2 globally for EDGE score) "
    "to island endemics like *Ilex dimorphophylla* from Japan's Ryukyu Islands."
)

body_para(doc,
    "In total, the FotW species with EDGE data collectively represent approximately 44,418 million years of "
    f"angiosperm evolutionary history — {ed_pct:.1f}% of all flowering plant ED on Earth. Among matched species, "
    "exceptional outliers stand out: *Gomortega keule* (78 Myr), the sole member of its family, survives in a "
    "single valley in Chile; and *Cephalotus follicularis* (76 Myr), a carnivorous pitcher plant from a single "
    "coastal habitat in Western Australia, represents one of the most isolated lineages in the plant kingdom."
)

section_heading(doc, "Extinct in the Wild: The Most Urgent Conservation Cases")

body_para(doc,
    f"Among the EDGE.List species documented in FotW, {len(ew_edge)} carry the most sobering designation on "
    "the IUCN Red List: Extinct in the Wild (EW). These are plants that no longer exist as self-sustaining "
    "populations in nature — their survival depends entirely on cultivation in botanical gardens, seed banks, "
    "and similar ex situ collections."
)

# EW species table
ew_tbl = doc.add_table(rows=1 + len(ew_edge_sorted), cols=5)
ew_tbl.style = "Table Grid"
headers = ["EDGE Rank", "Species", "Family", "EDGE Score", "ED (Myr)"]
header_row = ew_tbl.rows[0]
for j, h in enumerate(headers):
    hc = header_row.cells[j]
    set_cell_bg(hc, "7030A0")
    hp = hc.paragraphs[0]
    hr = hp.add_run(h)
    hr.bold = True
    hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hr.font.size = Pt(10)
    hp.paragraph_format.space_before = Pt(2)
    hp.paragraph_format.space_after  = Pt(2)

for i, row_data in enumerate(ew_edge_sorted):
    trow = ew_tbl.rows[i + 1]
    bg = "EAD1F5" if i % 2 == 0 else "F5EBFC"
    cells = trow.cells
    set_cell_bg(cells[0], bg)
    set_cell_bg(cells[1], bg)
    set_cell_bg(cells[2], bg)
    set_cell_bg(cells[3], bg)
    set_cell_bg(cells[4], bg)
    # Rank
    p0 = cells[0].paragraphs[0]
    p0.add_run(f"#{row_data['EDGE_rank']}").font.size = Pt(10)
    # Species with hyperlink
    p1 = cells[1].paragraphs[0]
    add_hyperlink(p1, row_data['species'].replace('_', ' '), row_data['taxon_url'])
    # Family
    p2 = cells[2].paragraphs[0]
    p2.add_run(row_data['family']).font.size = Pt(10)
    # EDGE score
    p3 = cells[3].paragraphs[0]
    p3.add_run(f"{float(row_data['edge_med']):.2f}").font.size = Pt(10)
    # ED
    p4 = cells[4].paragraphs[0]
    p4.add_run(f"{float(row_data['ed_med']):.1f}").font.size = Pt(10)
    for p in [p0, p1, p2, p3, p4]:
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

doc.add_paragraph()

body_para(doc,
    "The three *Brugmansia* species (angel's trumpets) represent a genus now considered entirely Extinct in the "
    "Wild. Once cultivated by Andean communities for centuries, these plants are no longer found in unmanaged "
    "natural populations. *Franklinia alatamaha*, a flowering tree from Georgia (USA), was last seen in the wild "
    "in 1803 and survives today solely through material preserved by early American botanists. FotW's "
    "documentation of these species links their occurrence records directly to the EDGE framework, making their "
    "conservation status visible to a global audience."
)

section_heading(doc, "The Davidson Endowed Chair and BSU's Conservation Mission")

body_para(doc,
    "Prof. Buerki holds the Davidson Endowed Chair of Botany at Boise State University — a position created "
    "to advance plant science and conservation at BSU and to build enduring partnerships with leading "
    "institutions worldwide. His involvement in the Forest et al. Science paper, and the subsequent "
    "implementation of EDGE scores in Flora of the World, is a direct expression of that mission."
)

body_para(doc,
    "Flora of the World is a global plant occurrence platform developed at the Royal Botanic Gardens, Kew. "
    "Through the Davidson Endowment, BSU is now a key contributor to the platform's evolution — specifically "
    "by embedding the EDGE framework into FotW's species pages. This will allow researchers, policymakers, "
    "and conservationists worldwide to see, alongside every specimen record, the evolutionary significance "
    "and extinction risk of each plant species."
)

body_para(doc,
    "The 229 EDGE.List species already documented in FotW represent a foundation for this integration. "
    "Expanding this coverage — and ensuring that EDGE scores and ED values are visible and accessible on "
    "the platform — is a concrete, ongoing goal of the work funded by the Davidson Endowment at BSU."
)

pull_quote(doc,
    "Our goal at Boise State is to make the results of this study directly visible in Flora of the World — "
    "so that anyone exploring a species page can see not just where a plant has been recorded, but how "
    "evolutionarily unique and how threatened it is. Embedding EDGE scores into the platform is central to "
    "the mission of the Davidson Endowment, and it is the work we are committed to completing.",
    "Prof. Sven Buerki, Davidson Endowed Chair of Botany, Boise State University"
)

body_para(doc,
    "As the Forest et al. study notes, protecting just 5.9% of species ranked by their EDGE score would "
    "safeguard half of all threatened angiosperm evolutionary history. BSU's role — through the Davidson "
    "Endowed Chair and the Flora of the World collaboration — in identifying where those species occur and "
    "making that information openly accessible, places the university at the heart of global plant "
    "conservation strategy."
)

doc.add_paragraph()

# ── TABLE: TOP 25 EDGE.LIST IN FOTW ──────────────────────────────────────────

section_heading(doc, "Top 25 EDGE.List Species Documented in Flora of the World")

p_note = doc.add_paragraph()
note_run = p_note.add_run(
    "Species names link directly to their Flora of the World taxon pages. "
    "Threat status follows the IUCN Red List. ED = Evolutionary Distinctiveness (millions of years). "
    "P(ext) = median probability of extinction."
)
note_run.italic = True
note_run.font.size = Pt(9)
note_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p_note.paragraph_format.space_after = Pt(6)

col_headers = ["Rank", "Species", "Family", "Threat", "EDGE\nScore", "ED\n(Myr)", "P(ext)", "FotW\nRecords"]
main_tbl = doc.add_table(rows=1 + len(top25), cols=len(col_headers))
main_tbl.style = "Table Grid"

# Header row
hrow = main_tbl.rows[0]
for j, h in enumerate(col_headers):
    hc = hrow.cells[j]
    set_cell_bg(hc, "1F4E79")
    hp = hc.paragraphs[0]
    hr = hp.add_run(h)
    hr.bold = True
    hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hr.font.size = Pt(9)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(2)
    hp.paragraph_format.space_after  = Pt(2)

for i, row_data in enumerate(top25):
    trow = main_tbl.rows[i + 1]
    threat = row_data["threat"]
    # alternating very light blue rows
    row_bg = "D9E2F3" if i % 2 == 0 else "EBF0FB"
    for cell in trow.cells:
        set_cell_bg(cell, row_bg)

    cells = trow.cells

    # Rank
    p = cells[0].paragraphs[0]
    p.add_run(f"#{row_data['EDGE_rank']}").font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Species — italic + hyperlink
    p1 = cells[1].paragraphs[0]
    url = row_data.get("taxon_url", "")
    sp_name = row_data["species"].replace("_", " ")
    if url:
        add_hyperlink(p1, sp_name, url)
    else:
        r = p1.add_run(sp_name)
        r.italic = True
        r.font.size = Pt(9)

    # Family
    p2 = cells[2].paragraphs[0]
    p2.add_run(row_data["family"]).font.size = Pt(9)

    # Threat — coloured badge
    p3 = cells[3].paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cells[3], threat_color(threat))
    tr = p3.add_run(threat)
    tr.bold = True
    tr.font.size = Pt(9)
    # white text for dark threat colours
    if threat in ("CR", "EN"):
        tr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    else:
        tr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # EDGE score
    p4 = cells[4].paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(f"{float(row_data['edge_med']):.2f}").font.size = Pt(9)

    # ED
    p5 = cells[5].paragraphs[0]
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.add_run(f"{float(row_data['ed_med']):.1f}").font.size = Pt(9)

    # P(ext)
    p6 = cells[6].paragraphs[0]
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pext_val = row_data.get("pext_med", "")
    try:
        p6.add_run(f"{float(pext_val):.3f}").font.size = Pt(9)
    except (ValueError, TypeError):
        p6.add_run("—").font.size = Pt(9)

    # FotW records
    p7 = cells[7].paragraphs[0]
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p7.add_run(row_data.get("n_fotw_records", "—")).font.size = Pt(9)

    for p in [p, p1, p2, p3, p4, p5, p6, p7]:
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

doc.add_paragraph()

# ── FOOTNOTE / SOURCE ────────────────────────────────────────────────────────
p_src = doc.add_paragraph()
p_src.paragraph_format.space_before = Pt(12)
src_run = p_src.add_run(
    "Source: Forest et al. (2026). EDGE scores for flowering plants. Science. "
    "EDGE gymnosperms precursor: Forest et al. (2018), Scientific Reports. "
    "FotW analysis by Prof. Sven Buerki (Boise State University, Davidson Endowed Chair of Botany) "
    "in collaboration with the Royal Botanic Gardens, Kew (Flora of the World project). "
    "Hyperlinks lead to live Flora of the World taxon pages."
)
src_run.font.size = Pt(9)
src_run.italic = True
src_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── SAVE ─────────────────────────────────────────────────────────────────────
out_path = "/Users/sven/Documents/Current_projects/Endowed_Chair/EDGE_flowering_plants/EDGE_FotW_press_article.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
